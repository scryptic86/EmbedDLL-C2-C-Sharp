from docx import Document
import re

INPUT_FILE  = "EmbedDLL_C2_Dropper-Loader_FINAL.docx"
OUTPUT_FILE = "EmbedDLL_C2_Dropper-Loader_FINAL_DEFANGED.docx"

# ---------------- Defanging logic ----------------

def defang_text(text):
    # Defang URL schemes
    text = re.sub(r'\bhttps://', 'hxxps://', text, flags=re.IGNORECASE)
    text = re.sub(r'\bhttp://',  'hxxp://',  text, flags=re.IGNORECASE)
    text = re.sub(r'\bftp://',   'fxp://',   text, flags=re.IGNORECASE)

    # Defang IPv4
    text = re.sub(
        r'\b(?:(?:25[0-5]|2[0-4]\d|1?\d?\d)\.){3}(?:25[0-5]|2[0-4]\d|1?\d?\d)\b',
        lambda m: m.group(0).replace('.', '[.]'),
        text
    )

    # Defang IPv6
    text = re.sub(
        r'\b([0-9a-fA-F]{1,4}:){2,7}[0-9a-fA-F]{1,4}\b',
        lambda m: m.group(0).replace(':', '[:]'),
        text
    )

    # Defang email addresses
    text = re.sub(
        r'\b([a-zA-Z0-9._%+-]+)@([a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b',
        lambda m: m.group(1) + '[@]' + m.group(2).replace('.', '[.]'),
        text
    )

    # Defang domains (avoid already defanged + common file extensions)
    def defang_domain(match):
        domain = match.group(0)

        if '[.]' in domain:
            return domain

        # Skip common file extensions and benign tokens
        if domain.lower().endswith((
            '.png','.jpg','.jpeg','.gif','.svg','.bmp','.webp',
            '.exe','.dll','.sys','.zip','.rar','.7z'
        )):
            return domain

        return domain.replace('.', '[.]')

    text = re.sub(
        r'\b(?!hxxps?://|fxp://)([a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}\b',
        defang_domain,
        text
    )

    return text

# ---------------- DOCX processing ----------------

def defang_paragraph(paragraph):
    for run in paragraph.runs:
        if run.text:
            run.text = defang_text(run.text)

def defang_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                defang_paragraph(paragraph)

def main():
    doc = Document(INPUT_FILE)

    # Defang normal paragraphs
    for paragraph in doc.paragraphs:
        defang_paragraph(paragraph)

    # Defang IOC tables (critical)
    for table in doc.tables:
        defang_table(table)

    doc.save(OUTPUT_FILE)

    print("\n[+] Defanging complete.")
    print(f"[+] Output written to: {OUTPUT_FILE}")
    print("[+] Formatting, images, and tables preserved.\n")

if __name__ == "__main__":
    main()
